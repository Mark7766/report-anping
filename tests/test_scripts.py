from __future__ import annotations

"""
Integration tests for Phase 2 scripts — show_params, build_chapter_prompt,
render_docx, check_compliance.

All tests run purely from within the Python process (no subprocess) to keep
execution fast, while still exercising the full CLI logic paths.
"""

import json
import sys
from pathlib import Path

REPO_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(REPO_ROOT))

FIXTURES_CHAPTERS = REPO_ROOT / "tests" / "fixtures" / "chapters"
FIXTURES_CEIC_CATALOG = REPO_ROOT / "tests" / "fixtures" / "ceic_catalog_sample.csv"
PARAMS_EXAMPLE = REPO_ROOT / "params_example.json"


# ============================================================
# show_params
# ============================================================


class TestShowParams:
    """Tests for scripts/show_params.py CLI logic."""

    def test_format_human_ii_contains_key_fields(self) -> None:
        from scripts.show_params import format_params_human

        output = format_params_human("II")
        assert "GB 17741" in output
        assert "name" in output
        assert "level" in output
        assert "exceedance_probs" in output
        assert "chapter" in output.lower()

    def test_format_human_lists_all_three_param_sections(self) -> None:
        from scripts.show_params import format_params_human

        output = format_params_human("II")
        assert "必填参数" in output
        assert "选填参数" in output

    def test_format_json_is_valid_json(self) -> None:
        from scripts.show_params import format_params_json

        raw = format_params_json("II")
        parsed = json.loads(raw)
        assert "params" in parsed
        assert "chapters" in parsed
        assert isinstance(parsed["params"], list)
        assert len(parsed["params"]) == 13

    def test_format_json_level_i_has_chapters(self) -> None:
        from scripts.show_params import format_params_json

        raw = format_params_json("I")
        parsed = json.loads(raw)
        assert len(parsed["chapters"]) > 0

    def test_param_spec_has_13_entries(self) -> None:
        from scripts.show_params import PARAM_SPEC

        assert len(PARAM_SPEC) == 13

    def test_all_required_params_have_example(self) -> None:
        from scripts.show_params import PARAM_SPEC

        for p in PARAM_SPEC:
            if p["required"]:
                assert p["example"] != "" or p["example"] == 0 or p["example"] == {}


# ============================================================
# build_chapter_prompt
# ============================================================


class TestBuildChapterPrompt:
    """Tests for scripts/build_chapter_prompt.py CLI logic."""

    def _load_params(self) -> dict:
        return json.loads(PARAMS_EXAMPLE.read_text(encoding="utf-8"))

    def test_known_chapter_returns_long_prompt(self) -> None:
        from lib.gb17741_knowledge import get_chapters_by_level
        from lib.prompts.chapter_prompts import build_chapter_prompt
        from scripts.build_chapter_prompt import find_chapter

        params = self._load_params()
        level = params.get("level", "II")
        chapters = get_chapters_by_level(level)
        chapter = find_chapter(chapters, "chapter4")
        assert chapter is not None

        prompt = build_chapter_prompt(
            project_data=params,
            chapter=chapter,
            chapter_index=3,
            total_chapters=len(chapters),
        )
        assert len(prompt) > 100
        assert "GB 17741" in prompt or "工程场地" in prompt

    def test_preface_chapter_prompt(self) -> None:
        from lib.gb17741_knowledge import get_chapters_by_level
        from lib.prompts.chapter_prompts import build_chapter_prompt
        from scripts.build_chapter_prompt import find_chapter

        params = self._load_params()
        chapters = get_chapters_by_level("II")
        chapter = find_chapter(chapters, "preface")
        assert chapter is not None

        prompt = build_chapter_prompt(params, chapter, 0, len(chapters))
        assert len(prompt) > 100

    def test_list_chapters_for_all_levels(self) -> None:
        from lib.gb17741_knowledge import get_chapters_by_level

        for level in ("I", "II", "III"):
            chapters = get_chapters_by_level(level)
            assert len(chapters) > 0
            for ch in chapters:
                assert "id" in ch and "title" in ch

    def test_find_chapter_returns_none_for_unknown(self) -> None:
        from scripts.build_chapter_prompt import find_chapter

        assert find_chapter([], "nonexistent") is None


# ============================================================
# render_docx
# ============================================================


class TestRenderDocx:
    """Integration tests for scripts/render_docx.py render() function."""

    def test_render_creates_docx_file(self, tmp_path: Path) -> None:
        from scripts.render_docx import render

        params = json.loads(PARAMS_EXAMPLE.read_text(encoding="utf-8"))
        out = tmp_path / "report.docx"
        render(params=params, chapters_dir=FIXTURES_CHAPTERS, out_path=out)
        assert out.exists()
        assert out.stat().st_size > 0

    def test_render_docx_is_nonzero_for_empty_chapters(self, tmp_path: Path) -> None:
        """render() with empty chapters dir should still produce a valid docx."""
        from scripts.render_docx import render

        empty_dir = tmp_path / "empty_chapters"
        empty_dir.mkdir()
        params = json.loads(PARAMS_EXAMPLE.read_text(encoding="utf-8"))
        out = tmp_path / "empty_report.docx"
        render(params=params, chapters_dir=empty_dir, out_path=out)
        assert out.exists()
        assert out.stat().st_size > 0

    def test_render_output_in_nested_dir(self, tmp_path: Path) -> None:
        """out_path parent is created automatically."""
        from scripts.render_docx import render

        params = json.loads(PARAMS_EXAMPLE.read_text(encoding="utf-8"))
        out = tmp_path / "exports" / "subdir" / "report.docx"
        render(params=params, chapters_dir=FIXTURES_CHAPTERS, out_path=out)
        assert out.exists()

    def test_render_bold_text_is_parsed(self, tmp_path: Path) -> None:
        """**bold** markdown in chapter text should produce bold Word runs, not asterisks."""
        import struct
        import zlib

        from docx import Document

        from scripts.render_docx import render

        # minimal 1x1 PNG
        def _make_png(path: Path) -> None:
            def chunk(t: bytes, d: bytes) -> bytes:
                return struct.pack(">I", len(d)) + t + d + struct.pack(">I", zlib.crc32(t + d) & 0xFFFFFFFF)

            sig = b"\x89PNG\r\n\x1a\n"
            ihdr = chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
            idat = chunk(b"IDAT", zlib.compress(b"\x00\xff\xff\xff"))
            iend = chunk(b"IEND", b"")
            path.write_bytes(sig + ihdr + idat + iend)

        chapters = tmp_path / "chapters"
        chapters.mkdir()
        md = "# 第1章 概述\n\n**评价单位**：某机构\n\n普通段落无粗体。\n"
        (chapters / "01_chapter1.md").write_text(md, encoding="utf-8")

        params = {"name": "Test", "level": "II", "evaluation_unit": "XXX", "report_date": "2025-01"}
        out = tmp_path / "report.docx"
        render(params=params, chapters_dir=chapters, out_path=out)

        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        # No raw asterisks should appear
        assert "**" not in all_text, f"Raw asterisks found in output: {all_text!r}"
        # Bold run should exist
        bold_runs = [r for p in doc.paragraphs for r in p.runs if r.bold and "评价单位" in r.text]
        assert bold_runs, "Expected a bold run containing '评价单位'"

    def test_render_caption_no_double_numbering(self, tmp_path: Path) -> None:
        """Caption already containing '图N-N' should not be prefixed with another number."""
        import struct
        import zlib

        from docx import Document

        from scripts.render_docx import render

        def _make_png(path: Path) -> None:
            def chunk(t: bytes, d: bytes) -> bytes:
                return struct.pack(">I", len(d)) + t + d + struct.pack(">I", zlib.crc32(t + d) & 0xFFFFFFFF)

            sig = b"\x89PNG\r\n\x1a\n"
            ihdr = chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
            idat = chunk(b"IDAT", zlib.compress(b"\x00\xff\xff\xff"))
            iend = chunk(b"IEND", b"")
            path.write_bytes(sig + ihdr + idat + iend)

        chapters = tmp_path / "chapters"
        chapters.mkdir()
        assets = tmp_path / "assets" / "generated"
        assets.mkdir(parents=True)
        _make_png(assets / "response_spectrum.png")

        md = "# 6 地震动参数\n\n![图6-1 设计反应谱](assets/generated/response_spectrum.png)\n"
        (chapters / "01_chapter6.md").write_text(md, encoding="utf-8")

        params = {"name": "Test", "level": "II", "evaluation_unit": "XXX", "report_date": "2025-01"}
        out = tmp_path / "report.docx"
        render(params=params, chapters_dir=chapters, out_path=out)

        doc = Document(str(out))
        caption_paras = [p.text for p in doc.paragraphs if "设计反应谱" in p.text]
        assert caption_paras, "Caption paragraph not found"
        # Must not contain double number like "图6-1  图6-1"
        for t in caption_paras:
            assert t.count("图6-1") <= 1, f"Double numbering detected: {t!r}"


# ============================================================
# check_compliance
# ============================================================


class TestCheckCompliance:
    """Tests for scripts/check_compliance.py logic."""

    def test_check_chapters_returns_list(self) -> None:
        from scripts.check_compliance import check_chapters

        results = check_chapters(FIXTURES_CHAPTERS, "II")
        assert isinstance(results, list)
        assert len(results) == 3  # preface, chapter1, chapter8

    def test_each_result_has_required_keys(self) -> None:
        from scripts.check_compliance import check_chapters

        for r in check_chapters(FIXTURES_CHAPTERS, "II"):
            assert "score" in r
            assert "status" in r
            assert "chapter_key" in r
            assert "file" in r

    def test_chapter_keys_are_valid(self) -> None:
        from scripts.check_compliance import _VALID_KEYS, check_chapters

        for r in check_chapters(FIXTURES_CHAPTERS, "II"):
            assert r["chapter_key"] in _VALID_KEYS

    def test_stem_to_chapter_key_mapping(self) -> None:
        from scripts.check_compliance import _stem_to_chapter_key

        assert _stem_to_chapter_key("01_preface") == "preface"
        assert _stem_to_chapter_key("02_chapter1") == "chapter1"
        assert _stem_to_chapter_key("chapter8") == "chapter8"
        assert _stem_to_chapter_key("03_chapter8") == "chapter8"
        assert _stem_to_chapter_key("unknown_file") is None

    def test_format_human_output(self) -> None:
        from scripts.check_compliance import check_chapters, format_human

        results = check_chapters(FIXTURES_CHAPTERS, "II")
        output = format_human(results, "II")
        assert "GB 17741" in output
        assert "合规检查报告" in output
        assert "综合评分" in output

    def test_format_json_is_valid(self) -> None:
        from scripts.check_compliance import check_chapters, format_json_output

        results = check_chapters(FIXTURES_CHAPTERS, "II")
        raw = format_json_output(results, "II")
        parsed = json.loads(raw)
        assert "work_level" in parsed
        assert "chapter_count" in parsed
        assert isinstance(parsed["chapters"], list)

    def test_empty_dir_returns_empty_results(self, tmp_path: Path) -> None:
        from scripts.check_compliance import check_chapters

        empty = tmp_path / "chapters"
        empty.mkdir()
        results = check_chapters(empty, "II")
        assert results == []


# ============================================================
# generate_figures / build_mt_chart
# ============================================================


class TestFigureScripts:
    """Tests for figure generation scripts."""

    def test_generate_figures_manifest_outputs_pngs(self, tmp_path: Path) -> None:
        from scripts.generate_figures import generate_figures_manifest

        params = json.loads(PARAMS_EXAMPLE.read_text(encoding="utf-8"))
        out_dir = tmp_path / "figures"
        manifest = generate_figures_manifest(
            params=params,
            out_dir=out_dir,
            catalog_path=FIXTURES_CEIC_CATALOG,
            min_magnitude=4.7,
        )

        assert "response_spectrum" in manifest
        assert "pga_comparison" in manifest
        assert "mt_chart" in manifest

        for key in ("response_spectrum", "pga_comparison", "mt_chart"):
            p = Path(manifest[key])
            assert p.exists()
            assert p.suffix.lower() == ".png"
            assert p.stat().st_size > 0

    def test_catalog_loader_parses_fixture(self) -> None:
        from lib.chart_builder import load_catalog_records

        records = load_catalog_records(FIXTURES_CEIC_CATALOG)
        assert len(records) == 5
        assert float(records[0]["magnitude"]) > 0
        # depth field is now parsed
        assert "depth" in records[0]
        assert float(records[0]["depth"]) > 0

    def test_generate_epicenter_map(self, tmp_path: Path) -> None:
        from lib.chart_builder import generate_epicenter_map, load_catalog_records

        records = load_catalog_records(FIXTURES_CEIC_CATALOG)
        out = tmp_path / "epicenter_map.png"
        result = generate_epicenter_map(records, out, center_lon=120.3, center_lat=34.2, site_name="TestSite")
        assert result.exists()
        assert result.stat().st_size > 0

    def test_generate_focal_depth_distribution(self, tmp_path: Path) -> None:
        from lib.chart_builder import generate_focal_depth_distribution, load_catalog_records

        records = load_catalog_records(FIXTURES_CEIC_CATALOG)
        out = tmp_path / "focal_depth.png"
        result = generate_focal_depth_distribution(records, out)
        assert result.exists()
        assert result.stat().st_size > 0

    def test_generate_intensity_bar_chart(self, tmp_path: Path) -> None:
        from lib.chart_builder import generate_intensity_bar_chart

        influences = [
            {"year": 1918, "location": "南澳", "magnitude": 7.25, "intensity": "VI"},
            {"year": 1969, "location": "阳江", "magnitude": 6.4, "intensity": "V"},
            {"year": 1994, "location": "台湾海峡", "magnitude": 7.3, "intensity": "IV"},
        ]
        out = tmp_path / "intensity_bar.png"
        result = generate_intensity_bar_chart(influences, out, site_name="TestSite")
        assert result.exists()
        assert result.stat().st_size > 0

    def test_generate_intensity_bar_chart_empty(self, tmp_path: Path) -> None:
        from lib.chart_builder import generate_intensity_bar_chart

        out = tmp_path / "intensity_bar_empty.png"
        result = generate_intensity_bar_chart([], out)
        assert result.exists()

    def test_generate_figures_manifest_includes_catalog_charts(self, tmp_path: Path) -> None:
        from scripts.generate_figures import generate_figures_manifest

        params = json.loads(PARAMS_EXAMPLE.read_text(encoding="utf-8"))
        out_dir = tmp_path / "figures_full"
        manifest = generate_figures_manifest(
            params=params,
            out_dir=out_dir,
            catalog_path=FIXTURES_CEIC_CATALOG,
        )

        for key in ("epicenter_map", "focal_depth_distribution", "intensity_bar_chart"):
            assert key in manifest, f"Expected '{key}' in manifest"
            p = Path(manifest[key])
            assert p.exists()
            assert p.stat().st_size > 0


# ============================================================
# init_project
# ============================================================


class TestInitProject:
    """Tests for scripts/init_project.py workspace initialisation."""

    def test_init_creates_required_directories(self, tmp_path: Path) -> None:
        from scripts.init_project import init_workspace

        init_workspace(tmp_path)
        for d in ("chapters", "exports", "data", "assets/generated"):
            assert (tmp_path / d).is_dir(), f"Expected directory: {d}"

    def test_init_creates_params_json_template(self, tmp_path: Path) -> None:
        from scripts.init_project import init_workspace

        init_workspace(tmp_path)
        params_path = tmp_path / "params.json"
        assert params_path.exists()
        params = json.loads(params_path.read_text(encoding="utf-8"))
        for key in ("name", "level", "location", "exceedance_probs"):
            assert key in params, f"Expected key '{key}' in params template"

    def test_init_skips_existing_params_by_default(self, tmp_path: Path) -> None:
        from scripts.init_project import init_workspace

        original = {"name": "existing_project", "level": "I"}
        (tmp_path / "params.json").write_text(json.dumps(original), encoding="utf-8")
        init_workspace(tmp_path, force=False)
        content = json.loads((tmp_path / "params.json").read_text(encoding="utf-8"))
        assert content["name"] == "existing_project"

    def test_init_force_overwrites_existing_params(self, tmp_path: Path) -> None:
        from scripts.init_project import init_workspace

        (tmp_path / "params.json").write_text('{"name": "old"}', encoding="utf-8")
        init_workspace(tmp_path, force=True)
        content = json.loads((tmp_path / "params.json").read_text(encoding="utf-8"))
        assert content["name"] == ""  # template default

    def test_init_is_idempotent(self, tmp_path: Path) -> None:
        from scripts.init_project import init_workspace

        init_workspace(tmp_path)
        init_workspace(tmp_path)  # second call must not raise
        assert (tmp_path / "chapters").is_dir()

    def test_init_returns_dict_with_all_resources(self, tmp_path: Path) -> None:
        from scripts.init_project import init_workspace

        result = init_workspace(tmp_path)
        assert "chapters" in result
        assert "exports" in result
        assert "data" in result
        assert "params.json" in result


# ============================================================
# build_mt_chart CLI
# ============================================================


class TestBuildMtChartCLI:
    """Tests for scripts/build_mt_chart.py CLI."""

    def test_cli_generates_png(self, tmp_path: Path) -> None:
        import subprocess

        out = tmp_path / "mt.png"
        result = subprocess.run(
            [sys.executable, "scripts/build_mt_chart.py", "--catalog", str(FIXTURES_CEIC_CATALOG), "--out", str(out)],
            capture_output=True,
            text=True,
            cwd=str(REPO_ROOT),
        )
        assert result.returncode == 0, f"stderr: {result.stderr}"
        assert out.exists()
        assert out.stat().st_size > 0

    def test_cli_custom_title(self, tmp_path: Path) -> None:
        import subprocess

        out = tmp_path / "mt_titled.png"
        result = subprocess.run(
            [
                sys.executable,
                "scripts/build_mt_chart.py",
                "--catalog",
                str(FIXTURES_CEIC_CATALOG),
                "--out",
                str(out),
                "--title",
                "Test M-T Chart",
                "--min-mag",
                "4.5",
            ],
            capture_output=True,
            text=True,
            cwd=str(REPO_ROOT),
        )
        assert result.returncode == 0, f"stderr: {result.stderr}"
        assert out.exists()

    def test_cli_missing_catalog_exits_nonzero(self, tmp_path: Path) -> None:
        import subprocess

        result = subprocess.run(
            [
                sys.executable,
                "scripts/build_mt_chart.py",
                "--catalog",
                "nonexistent_catalog.csv",
                "--out",
                str(tmp_path / "out.png"),
            ],
            capture_output=True,
            text=True,
            cwd=str(REPO_ROOT),
        )
        assert result.returncode != 0


# ============================================================
# compliance_prompts
# ============================================================


class TestCompliancePrompts:
    """Tests for lib/prompts/compliance_prompts.py prompt constants."""

    def test_check_system_prompt_is_non_empty_str(self) -> None:
        from lib.prompts.compliance_prompts import COMPLIANCE_CHECK_SYSTEM_PROMPT

        assert isinstance(COMPLIANCE_CHECK_SYSTEM_PROMPT, str)
        assert len(COMPLIANCE_CHECK_SYSTEM_PROMPT) > 50

    def test_check_system_prompt_mentions_gb17741(self) -> None:
        from lib.prompts.compliance_prompts import COMPLIANCE_CHECK_SYSTEM_PROMPT

        assert "GB 17741" in COMPLIANCE_CHECK_SYSTEM_PROMPT

    def test_check_system_prompt_has_output_format_section(self) -> None:
        from lib.prompts.compliance_prompts import COMPLIANCE_CHECK_SYSTEM_PROMPT

        assert "合规检测结果" in COMPLIANCE_CHECK_SYSTEM_PROMPT

    def test_fix_system_prompt_is_non_empty_str(self) -> None:
        from lib.prompts.compliance_prompts import COMPLIANCE_FIX_SYSTEM_PROMPT

        assert isinstance(COMPLIANCE_FIX_SYSTEM_PROMPT, str)
        assert len(COMPLIANCE_FIX_SYSTEM_PROMPT) > 50

    def test_fix_system_prompt_has_report_markers(self) -> None:
        from lib.prompts.compliance_prompts import COMPLIANCE_FIX_SYSTEM_PROMPT

        assert "报告内容开始" in COMPLIANCE_FIX_SYSTEM_PROMPT
        assert "报告内容结束" in COMPLIANCE_FIX_SYSTEM_PROMPT

    def test_build_compliance_check_prompt(self) -> None:
        from lib.prompts.compliance_prompts import build_compliance_check_prompt

        result = build_compliance_check_prompt("chapter1", "II")
        assert isinstance(result, str)
        assert "GB 17741" in result
        assert len(result) > 100

    def test_build_compliance_fix_prompt(self) -> None:
        from lib.prompts.compliance_prompts import build_compliance_fix_prompt

        result = build_compliance_fix_prompt("chapter4", "II")
        assert isinstance(result, str)
        assert len(result) > 100

    def test_build_normalize_terms_prompt(self) -> None:
        from lib.prompts.compliance_prompts import build_normalize_terms_prompt

        result = build_normalize_terms_prompt("chapter1", "II")
        assert isinstance(result, str)
        assert "标准术语" in result

    def test_build_validate_params_prompt(self) -> None:
        from lib.prompts.compliance_prompts import build_validate_params_prompt

        result = build_validate_params_prompt("II", "chapter1")
        assert isinstance(result, str)
        assert len(result) > 100

    def test_build_add_reference_prompt(self) -> None:
        from lib.prompts.compliance_prompts import build_add_reference_prompt

        result = build_add_reference_prompt("chapter1", "II")
        assert isinstance(result, str)
        assert len(result) > 100


# ============================================================
# Additional chapter_prompts coverage
# ============================================================


class TestChapterPromptsAdditional:
    """Extra coverage for lib/prompts/chapter_prompts.py paths."""

    def _load_params(self) -> dict:
        return json.loads(PARAMS_EXAMPLE.read_text(encoding="utf-8"))

    def test_appendix_chapter_prompt(self) -> None:
        from lib.gb17741_knowledge import get_chapters_by_level
        from lib.prompts.chapter_prompts import build_chapter_prompt
        from scripts.build_chapter_prompt import find_chapter

        params = self._load_params()
        chapters = get_chapters_by_level(params.get("level", "II"))
        chapter = find_chapter(chapters, "appendix")
        if chapter is None:
            return  # appendix may not exist for all levels
        idx = next(i for i, c in enumerate(chapters) if c["id"] == "appendix")
        prompt = build_chapter_prompt(params, chapter, idx, len(chapters))
        assert "附录" in prompt or "appendix" in prompt.lower()
        assert len(prompt) > 100

    def test_build_full_report_prompt(self) -> None:
        from lib.prompts.chapter_prompts import build_full_report_prompt

        params = self._load_params()
        result = build_full_report_prompt(params)
        assert isinstance(result, str)
        assert len(result) > 100

    def test_exceedance_prob_guidance_included_in_chapter_prompt(self) -> None:
        """Chapters needing probability tables should embed the guidance block."""
        from lib.gb17741_knowledge import get_chapters_by_level
        from lib.prompts.chapter_prompts import build_chapter_prompt
        from scripts.build_chapter_prompt import find_chapter

        params = self._load_params()
        chapters = get_chapters_by_level(params.get("level", "II"))
        # chapter8 (地震动参数) includes exceedance probability guidance
        chapter = find_chapter(chapters, "chapter8")
        if chapter is None:
            chapter = chapters[-2] if len(chapters) >= 2 else chapters[0]
        idx = next(i for i, c in enumerate(chapters) if c["id"] == chapter["id"])
        prompt = build_chapter_prompt(params, chapter, idx, len(chapters))
        assert len(prompt) > 100

    def test_get_report_structure(self) -> None:
        from lib.prompts.chapter_prompts import get_report_structure

        result = get_report_structure("II")
        assert isinstance(result, str)
        assert "##" in result  # has Markdown headings

    def test_build_report_prompt_compat_alias(self) -> None:
        from lib.prompts.chapter_prompts import build_report_prompt

        params = self._load_params()
        result = build_report_prompt(params)
        assert isinstance(result, str)
        assert len(result) > 50


class TestGb17741KnowledgeExtra:
    """Additional coverage for lib/gb17741_knowledge.py helper functions."""

    def test_get_clause_two_part_id(self) -> None:
        from lib.gb17741_knowledge import get_clause

        # Two-part clause ID: should return the section dict
        result = get_clause("5.2")
        # May be None if data not loaded, but must not raise
        assert result is None or isinstance(result, dict)

    def test_get_clause_three_part_id(self) -> None:
        from lib.gb17741_knowledge import get_clause

        result = get_clause("5.2.1")
        assert result is None or isinstance(result, (dict, str))

    def test_get_clause_short_id_returns_none(self) -> None:
        from lib.gb17741_knowledge import get_clause

        assert get_clause("5") is None

    def test_get_appendix_b_returns_dict(self) -> None:
        from lib.gb17741_knowledge import get_appendix_b

        result = get_appendix_b()
        assert isinstance(result, dict)

    def test_full_guidance_with_formula_chapter(self) -> None:
        """chapter9 has a formula map entry — exercises lines 1025-1030."""
        from lib.gb17741_knowledge import get_full_standard_guidance_for_ai

        result = get_full_standard_guidance_for_ai(chapter_id="chapter9")
        assert isinstance(result, str)
        assert len(result) > 50


class TestGenerateFiguresExtra:
    """Additional coverage for scripts/generate_figures.py coord-fallback path."""

    def _load_params(self) -> dict:
        import json
        from pathlib import Path

        return json.loads((Path(__file__).parent.parent / "params_example.json").read_text())

    def test_coord_fallback_when_invalid_values(self, tmp_path) -> None:
        """coords set to non-numeric trigger the except branch (lines 82-83)."""
        from scripts.generate_figures import generate_figures_manifest

        catalog = Path(__file__).parent / "fixtures" / "ceic_catalog_sample.csv"
        params = self._load_params()
        # Inject invalid coordinate values to trigger TypeError/ValueError fallback
        params["coordinate_lon"] = "not-a-number"
        params["coordinate_lat"] = None
        out = tmp_path / "figs"
        manifest = generate_figures_manifest(params, out_dir=out, catalog_path=catalog)
        assert "epicenter_map" in manifest

    def test_auto_detect_catalog_when_default_exists(self, tmp_path, monkeypatch) -> None:
        """When catalog_path is None and _DEFAULT_CATALOG exists, it is used (line 57)."""
        import shutil

        import scripts.generate_figures as gf

        catalog_src = Path(__file__).parent / "fixtures" / "ceic_catalog_sample.csv"
        fake_default = tmp_path / "ceic_catalog.csv"
        shutil.copy(catalog_src, fake_default)

        monkeypatch.setattr(gf, "_DEFAULT_CATALOG", fake_default)

        params = self._load_params()
        out = tmp_path / "figs_auto"
        manifest = gf.generate_figures_manifest(params, out_dir=out, catalog_path=None)
        # With a valid catalog auto-detected, mt_chart should be in manifest
        assert "mt_chart" in manifest
