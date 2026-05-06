# -*- coding: utf-8 -*-
"""
ajepro backend - 专业表格渲染引擎
将Markdown表格转换为专业的Word表格（三线表/全框表样式）
"""

import re

from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# 使用数值常量代替枚举
WD_ALIGN_PARAGRAPH_CENTER = 1
WD_ALIGN_PARAGRAPH_LEFT = 0
WD_ALIGN_PARAGRAPH_RIGHT = 2

from lib.logger import get_logger

logger = get_logger("table_renderer")


class MarkdownTableParser:
    """Markdown表格解析器"""

    @staticmethod
    def parse(lines):
        """
        解析Markdown表格行为结构化数据

        Args:
            lines: 表格行列表（包含 | 分隔的行）

        Returns:
            dict: {
                'headers': ['列1', '列2', ...],
                'alignments': ['left', 'center', 'right', ...],
                'rows': [['数据1', '数据2', ...], ...]
            }
        """
        if not lines:
            return None

        result = {"headers": [], "alignments": [], "rows": []}

        header_parsed = False
        separator_parsed = False

        for line in lines:
            line = line.strip()
            if not line.startswith("|"):
                continue

            # 解析单元格
            cells = MarkdownTableParser._parse_cells(line)

            if not cells:
                continue

            # 第一行是表头
            if not header_parsed:
                result["headers"] = cells
                header_parsed = True
                continue

            # 第二行是对齐分隔符
            if not separator_parsed:
                result["alignments"] = MarkdownTableParser._parse_alignments(cells)
                separator_parsed = True
                continue

            # 后续是数据行
            result["rows"].append(cells)

        return result if result["headers"] else None

    @staticmethod
    def _parse_cells(line):
        """解析表格行中的单元格"""
        # 去掉首尾的 |
        line = line.strip()
        if line.startswith("|"):
            line = line[1:]
        if line.endswith("|"):
            line = line[:-1]

        # 分割单元格
        cells = [cell.strip() for cell in line.split("|")]
        return cells

    @staticmethod
    def _parse_alignments(cells):
        """解析对齐方式"""
        alignments = []
        for cell in cells:
            cell = cell.strip()
            if cell.startswith(":") and cell.endswith(":"):
                alignments.append("center")
            elif cell.endswith(":"):
                alignments.append("right")
            elif cell.startswith(":"):
                alignments.append("left")
            else:
                alignments.append("left")  # 默认左对齐
        return alignments

    @staticmethod
    def is_numeric(value):
        """判断值是否为数字（用于自动右对齐）"""
        if not value:
            return False
        # 去掉常见的数字格式符号
        cleaned = value.replace(",", "").replace("%", "").replace("°", "").strip()
        try:
            float(cleaned)
            return True
        except ValueError:
            # 检查是否是范围值如 "0.05~0.10"
            if "~" in cleaned or "-" in cleaned:
                parts = re.split(r"[~\-]", cleaned)
                return all(MarkdownTableParser.is_numeric(p) for p in parts if p)
            return False


class DocxTableRenderer:
    """Word表格渲染器"""

    # 表格样式常量
    STYLE_THREE_LINE = "three-line"  # 三线表
    STYLE_FULL_BORDER = "full-border"  # 全框表
    STYLE_SIMPLE = "simple"  # 简约表（无边框）

    # 边框粗细（单位：1/8磅）
    BORDER_THICK = 12  # 1.5pt
    BORDER_THIN = 4  # 0.5pt

    def __init__(self, doc, format_settings=None):
        """
        初始化表格渲染器

        Args:
            doc: Word文档对象
            format_settings: 格式设置字典
        """
        self.doc = doc
        self.format_settings = format_settings or {}

    def render(self, table_data, caption=None, table_number=None, style=None):
        """
        渲染表格到Word文档

        Args:
            table_data: 解析后的表格数据
            caption: 表格标题
            table_number: 表格编号（如 "3-1"）
            style: 表格样式（three-line/full-border/simple）

        Returns:
            table: Word表格对象
        """
        if not table_data or not table_data.get("headers"):
            return None

        # 🔑 修复：从格式设置读取表格样式，确保正确应用
        configured_style = self.format_settings.get("table_style", self.STYLE_FULL_BORDER)
        style = style or configured_style

        # 记录使用的样式
        logger.info(f"[TableRenderer] 表格样式: 配置={configured_style}, 最终使用={style}")

        headers = table_data["headers"]
        rows = table_data.get("rows", [])
        alignments = table_data.get("alignments", ["left"] * len(headers))

        # 添加表题
        if caption or table_number:
            self._add_table_caption(caption, table_number)

        # 创建表格
        row_count = len(rows) + 1  # +1 for header
        col_count = len(headers)

        table = self.doc.add_table(rows=row_count, cols=col_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 设置表格宽度为页面宽度
        self._set_table_width(table)

        # 填充表头
        header_row = table.rows[0]
        for i, header_text in enumerate(headers):
            cell = header_row.cells[i]
            self._set_cell_content(cell, header_text, is_header=True)
            self._set_cell_alignment(cell, alignments[i] if i < len(alignments) else "left")

        # 填充数据行
        for row_idx, row_data in enumerate(rows):
            table_row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < col_count:
                    cell = table_row.cells[col_idx]
                    self._set_cell_content(cell, cell_text, is_header=False)

                    # 智能对齐：数字自动右对齐
                    if MarkdownTableParser.is_numeric(cell_text):
                        self._set_cell_alignment(cell, "right")
                    elif col_idx < len(alignments):
                        self._set_cell_alignment(cell, alignments[col_idx])

        # 应用边框样式
        if style == self.STYLE_THREE_LINE:
            self._apply_three_line_style(table)
        elif style == self.STYLE_FULL_BORDER:
            self._apply_full_border_style(table)
        else:
            self._apply_simple_style(table)

        # 添加表格后的间距
        self.doc.add_paragraph()

        return table

    def _add_table_caption(self, caption, table_number):
        """添加表题"""
        caption_font = self.format_settings.get("table_caption_font", "黑体")
        caption_size = self.format_settings.get("table_caption_size", 10.5)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH_CENTER

        # 构建表题文本
        if table_number:
            caption_text = f"表{table_number}"
            if caption:
                caption_text += f"  {caption}"
        else:
            caption_text = caption or ""

        run = p.add_run(caption_text)
        run.font.name = caption_font
        run.font.size = Pt(caption_size)
        run.font.bold = True
        run.font.italic = False  # 🔑 确保不使用斜体
        from docx.shared import RGBColor

        run.font.color.rgb = RGBColor(0, 0, 0)  # 🔑 显式设置为黑色
        run._element.rPr.rFonts.set(qn("w:eastAsia"), caption_font)

        # 设置段后间距
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.space_before = Pt(6)

    def _set_table_width(self, table):
        """设置表格宽度为页面可用宽度"""
        # 获取页边距
        margin_left = self.format_settings.get("page_margin_left", 3.0)
        margin_right = self.format_settings.get("page_margin_right", 2.5)

        # A4纸宽度21cm，减去左右页边距
        available_width = 21.0 - margin_left - margin_right

        # 设置表格宽度
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), str(int(available_width * 567)))  # 转换为twips
        tblW.set(qn("w:type"), "dxa")
        tblPr.append(tblW)

        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def _set_cell_content(self, cell, text, is_header=False):
        """设置单元格内容"""
        content_font = self.format_settings.get("table_content_font", "宋体")
        content_size = self.format_settings.get("table_content_size", 10.5)

        # 清空单元格
        cell.text = ""

        # 添加内容
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = content_font
        run.font.size = Pt(content_size)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), content_font)

        if is_header:
            run.font.bold = True

        # 设置单元格垂直居中
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # 设置单元格内边距
        self._set_cell_margins(cell)

    def _set_cell_margins(self, cell):
        """设置单元格内边距"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        tcMar = OxmlElement("w:tcMar")
        for margin_name in ["top", "left", "bottom", "right"]:
            margin = OxmlElement(f"w:{margin_name}")
            margin.set(qn("w:w"), "57")  # 约0.1cm
            margin.set(qn("w:type"), "dxa")
            tcMar.append(margin)

        tcPr.append(tcMar)

    def _set_cell_alignment(self, cell, alignment):
        """设置单元格对齐方式"""
        p = cell.paragraphs[0]
        if alignment == "center":
            p.alignment = WD_ALIGN_PARAGRAPH_CENTER
        elif alignment == "right":
            p.alignment = WD_ALIGN_PARAGRAPH_RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH_LEFT

    def _apply_three_line_style(self, table):
        """应用三线表样式 - 加强版"""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

        # 创建表格边框设置
        tblBorders = OxmlElement("w:tblBorders")

        # 顶线：粗线（1.5pt）
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), str(self.BORDER_THICK))  # 12 = 1.5pt
        top.set(qn("w:color"), "000000")
        top.set(qn("w:space"), "0")
        tblBorders.append(top)

        # 底线：粗线（1.5pt）
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(self.BORDER_THICK))  # 12 = 1.5pt
        bottom.set(qn("w:color"), "000000")
        bottom.set(qn("w:space"), "0")
        tblBorders.append(bottom)

        # 清除左右边框和内部垂直线
        for border_name in ["left", "right", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            border.set(qn("w:sz"), "0")
            tblBorders.append(border)

        # 清除内部水平线（除了表头下方的栏目线）
        insideH = OxmlElement("w:insideH")
        insideH.set(qn("w:val"), "none")
        insideH.set(qn("w:sz"), "0")
        tblBorders.append(insideH)

        tblPr.append(tblBorders)

        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

        # 为表头行添加底部边框（栏目线 - 细线 0.5pt）
        if len(table.rows) > 0:
            header_row = table.rows[0]
            for cell in header_row.cells:
                self._set_cell_bottom_border(cell, self.BORDER_THIN)  # 4 = 0.5pt

        # 🔑 关键：设置表格布局为自动调整
        tblLayout = OxmlElement("w:tblLayout")
        tblLayout.set(qn("w:type"), "autofit")
        tblPr.append(tblLayout)

    def _apply_full_border_style(self, table):
        """应用全框表样式"""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

        # 创建表格边框设置
        tblBorders = OxmlElement("w:tblBorders")

        # 所有边框：细线
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), str(self.BORDER_THIN))
            border.set(qn("w:color"), "000000")
            tblBorders.append(border)

        tblPr.append(tblBorders)

        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def _apply_simple_style(self, table):
        """应用简约表样式（无边框）"""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

        # 创建表格边框设置
        tblBorders = OxmlElement("w:tblBorders")

        # 清除所有边框
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")
            tblBorders.append(border)

        tblPr.append(tblBorders)

        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def _set_cell_bottom_border(self, cell, size):
        """设置单元格底部边框"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        tcBorders = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(size))
        bottom.set(qn("w:color"), "000000")
        tcBorders.append(bottom)

        tcPr.append(tcBorders)


class TableNumbering:
    """表格编号追踪器"""

    def __init__(self):
        self.chapter = 0
        self.table_count = 0
        self.table_registry = {}  # id -> number mapping

    def set_chapter(self, chapter_num):
        """设置当前章节号（支持数字和字符串，如"前"、"附A"）"""
        # 🔑 支持字符串章节号（如"前"、"附A"）
        if isinstance(chapter_num, str):
            # 字符串章节号，直接使用
            if chapter_num != self.chapter:
                self.chapter = chapter_num
                self.table_count = 0
        else:
            # 数字章节号，进行保护性检查
            if chapter_num <= 0:
                chapter_num = 4
                logger.warning(f"[TableNumbering] 检测到无效章节号，使用默认值: {chapter_num}")

            if chapter_num != self.chapter:
                self.chapter = chapter_num
                self.table_count = 0

    def next_number(self, table_id=None):
        """
        获取下一个表格编号

        Args:
            table_id: 可选的表格ID用于交叉引用

        Returns:
            str: 表格编号，如 "3-1", "前-1", "附A-1"
        """
        # 🔑 保护性检查：确保章节号有效（仅对数字章节）
        if isinstance(self.chapter, int) and self.chapter <= 0:
            self.chapter = 4
            logger.warning(f"[TableNumbering] 表格编号时检测到无效章节号，使用默认值: {self.chapter}")

        self.table_count += 1
        number = f"{self.chapter}-{self.table_count}"

        if table_id:
            self.table_registry[table_id] = number

        return number

    def get_number(self, table_id):
        """根据ID获取表格编号"""
        return self.table_registry.get(table_id)
