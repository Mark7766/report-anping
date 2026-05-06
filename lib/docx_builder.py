# -*- coding: utf-8 -*-
"""
ajepro backend - Word文档构建器工具
封装python-docx底层XML操作，处理目录、页眉页脚等高级功能
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

# 使用数值常量代替枚举，避免兼容性问题
WD_ALIGN_PARAGRAPH_CENTER = 1
WD_ALIGN_PARAGRAPH_LEFT = 0
WD_ALIGN_PARAGRAPH_RIGHT = 2

from lib.logger import get_logger

logger = get_logger("docx_builder")


class DocxBuilder:
    """Word文档构建器 - 封装底层XML操作"""

    def __init__(self, doc=None):
        """
        初始化构建器

        Args:
            doc: 现有Word文档对象，如果为None则创建新文档
        """
        self.doc = doc or Document()

    def get_document(self):
        """获取文档对象"""
        return self.doc

    # ==================== 目录相关 ====================

    def add_toc(self, title="目  录", levels=3, show_page_numbers=True, use_hyperlinks=True):
        """
        添加真正可更新的Word目录

        Args:
            title: 目录标题
            levels: 目录级别 (1-9)
            show_page_numbers: 是否显示页码
            use_hyperlinks: 是否使用超链接
        """
        # 添加目录标题
        toc_heading = self.doc.add_paragraph()
        toc_heading.alignment = WD_ALIGN_PARAGRAPH_CENTER
        run = toc_heading.add_run(title)
        run.font.name = "黑体"
        run.font.size = Pt(18)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

        # 设置段后间距
        toc_heading.paragraph_format.space_after = Pt(24)

        # 创建目录段落
        paragraph = self.doc.add_paragraph()

        # 构建TOC域代码
        # TOC \o "1-3" \h \z \u
        # \o "1-3" : 包含1-3级标题
        # \h : 使用超链接
        # \z : 隐藏Web视图中的制表符前导符和页码
        # \u : 使用应用的段落大纲级别
        toc_options = f'\\o "1-{levels}"'
        if use_hyperlinks:
            toc_options += " \\h"
        if show_page_numbers:
            toc_options += " \\z"
        toc_options += " \\u"

        self._add_field(paragraph, f"TOC {toc_options}", '（请右键点击此处选择"更新域"，或按 F9 更新目录）')

        # 添加分页符
        self.doc.add_page_break()

        logger.info("已添加可更新的Word目录")

    def _add_field(self, paragraph, field_code, placeholder_text=""):
        """
        向段落添加Word域代码

        Args:
            paragraph: 段落对象
            field_code: 域代码（如 "TOC \\o '1-3'"）
            placeholder_text: 占位文本
        """
        run = paragraph.add_run()
        r = run._r

        # fldChar begin
        fldChar_begin = OxmlElement("w:fldChar")
        fldChar_begin.set(qn("w:fldCharType"), "begin")
        r.append(fldChar_begin)

        # instrText - 域代码
        run2 = paragraph.add_run()
        r2 = run2._r
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = field_code
        r2.append(instrText)

        # fldChar separate
        run3 = paragraph.add_run()
        r3 = run3._r
        fldChar_sep = OxmlElement("w:fldChar")
        fldChar_sep.set(qn("w:fldCharType"), "separate")
        r3.append(fldChar_sep)

        # 占位文本
        if placeholder_text:
            run4 = paragraph.add_run(placeholder_text)
            run4.font.name = "宋体"
            run4.font.size = Pt(10)
            run4.font.italic = True
            run4._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # fldChar end
        run5 = paragraph.add_run()
        r5 = run5._r
        fldChar_end = OxmlElement("w:fldChar")
        fldChar_end.set(qn("w:fldCharType"), "end")
        r5.append(fldChar_end)

    # ==================== 页眉页脚相关 ====================

    def setup_headers_footers(
        self,
        header_odd=None,
        header_even=None,
        footer_style="- X -",
        first_page_different=True,
        odd_even_different=True,
        header_font="宋体",
        header_size=9,
        show_header_line=True,
    ):
        """
        设置页眉页脚

        Args:
            header_odd: 奇数页页眉内容
            header_even: 偶数页页眉内容
            footer_style: 页脚样式 ('- X -', 'X', '第X页', '第X页 共Y页')
            first_page_different: 首页是否不同（封面无页眉页脚）
            odd_even_different: 奇偶页是否不同
            header_font: 页眉字体
            header_size: 页眉字号
            show_header_line: 是否显示页眉下划线
        """
        section = self.doc.sections[0]

        # 设置首页不同
        section.different_first_page_header_footer = first_page_different

        # 设置奇偶页不同
        self.doc.settings.odd_and_even_pages_header_footer = odd_even_different

        # 设置页眉
        if header_odd or header_even:
            self._setup_header(
                section, header_odd, header_even, header_font, header_size, show_header_line, odd_even_different
            )

        # 设置页脚（页码）
        self._setup_footer(section, footer_style, first_page_different)

        logger.info("已设置页眉页脚")

    def _setup_header(self, section, header_odd, header_even, font_name, font_size, show_line, odd_even_different):
        """设置页眉"""
        # 奇数页页眉（或统一页眉）
        header = section.header
        header.is_linked_to_previous = False

        if header_odd:
            self._clear_header_footer(header)
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH_RIGHT
            run = p.add_run(header_odd)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

            if show_line:
                self._add_header_bottom_border(p)

        # 偶数页页眉
        if odd_even_different and header_even:
            even_header = section.even_page_header
            even_header.is_linked_to_previous = False

            self._clear_header_footer(even_header)
            p = even_header.paragraphs[0] if even_header.paragraphs else even_header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH_LEFT
            run = p.add_run(header_even)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

            if show_line:
                self._add_header_bottom_border(p)

    def _setup_footer(self, section, footer_style, first_page_different):
        """设置页脚（页码）"""
        footer = section.footer
        footer.is_linked_to_previous = False

        self._clear_header_footer(footer)
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH_CENTER

        # 添加页码域
        self._add_page_number_field(p, footer_style)

        # 偶数页页脚（保持相同的页码格式）
        even_footer = section.even_page_footer
        even_footer.is_linked_to_previous = False

        self._clear_header_footer(even_footer)
        p2 = even_footer.paragraphs[0] if even_footer.paragraphs else even_footer.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH_CENTER
        self._add_page_number_field(p2, footer_style)

    def _add_page_number_field(self, paragraph, style="- X -"):
        """添加页码域"""
        run = paragraph.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

        if style == "- X -":
            # 格式: - 1 -
            paragraph.add_run("- ")
            self._add_field(paragraph, "PAGE", "")
            paragraph.add_run(" -")
        elif style == "X":
            # 格式: 1
            self._add_field(paragraph, "PAGE", "")
        elif style == "第X页":
            # 格式: 第 1 页
            paragraph.add_run("第 ")
            self._add_field(paragraph, "PAGE", "")
            paragraph.add_run(" 页")
        elif style == "第X页 共Y页":
            # 格式: 第 1 页 共 10 页
            paragraph.add_run("第 ")
            self._add_field(paragraph, "PAGE", "")
            paragraph.add_run(" 页 共 ")
            self._add_field(paragraph, "NUMPAGES", "")
            paragraph.add_run(" 页")
        else:
            # 默认格式
            self._add_field(paragraph, "PAGE", "")

    def _add_header_bottom_border(self, paragraph):
        """添加页眉底部边框线"""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")

        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")  # 0.75pt
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")

        pBdr.append(bottom)
        pPr.append(pBdr)

    def _clear_header_footer(self, header_footer):
        """清空页眉或页脚内容"""
        for p in header_footer.paragraphs:
            p.clear()

    # ==================== 页面设置相关 ====================

    def setup_page(
        self,
        paper_size="A4",
        orientation="portrait",
        margin_top=2.54,
        margin_bottom=2.54,
        margin_left=3.0,
        margin_right=2.5,
        gutter=0,
        gutter_position="left",
    ):
        """
        设置页面属性

        Args:
            paper_size: 纸张大小 ('A4', 'Letter', 'Legal')
            orientation: 方向 ('portrait', 'landscape')
            margin_top: 上边距 (cm)
            margin_bottom: 下边距 (cm)
            margin_left: 左边距 (cm)
            margin_right: 右边距 (cm)
            gutter: 装订线宽度 (cm)
            gutter_position: 装订线位置 ('left', 'top')
        """
        section = self.doc.sections[0]

        # 设置纸张大小
        if paper_size == "A4":
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
        elif paper_size == "Letter":
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        elif paper_size == "Legal":
            section.page_width = Inches(8.5)
            section.page_height = Inches(14)

        # 设置方向
        if orientation == "landscape":
            # WD_ORIENT.LANDSCAPE = 1
            from docx.enum.section import WD_ORIENT

            section.orientation = WD_ORIENT.LANDSCAPE
            # 交换宽高
            section.page_width, section.page_height = section.page_height, section.page_width

        # 设置页边距
        section.top_margin = Cm(margin_top)
        section.bottom_margin = Cm(margin_bottom)
        section.left_margin = Cm(margin_left)
        section.right_margin = Cm(margin_right)

        # 设置装订线
        if gutter > 0:
            section.gutter = Cm(gutter)

        logger.info(f"页面设置完成: {paper_size}, 边距({margin_top}, {margin_bottom}, {margin_left}, {margin_right})cm")

    # ==================== 分节符相关 ====================

    def add_section_break(self, break_type="next_page"):
        """
        添加分节符

        Args:
            break_type: 分节符类型 ('next_page', 'continuous', 'odd_page', 'even_page')
        """
        from docx.enum.section import WD_SECTION_START

        type_map = {
            "next_page": WD_SECTION_START.NEW_PAGE,
            "continuous": WD_SECTION_START.CONTINUOUS,
            "odd_page": WD_SECTION_START.ODD_PAGE,
            "even_page": WD_SECTION_START.EVEN_PAGE,
        }

        new_section = self.doc.add_section(type_map.get(break_type, WD_SECTION_START.NEW_PAGE))
        return new_section

    # ==================== 样式设置相关 ====================

    def setup_toc_styles(self):
        """设置目录样式"""
        styles = self.doc.styles

        # TOC 1 - 一级目录
        try:
            toc1 = styles["TOC 1"]
        except KeyError:
            toc1 = styles.add_style("TOC 1", 1)  # 1 = paragraph

        toc1.font.name = "黑体"
        toc1.font.size = Pt(14)
        toc1.font.bold = False
        toc1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        toc1.paragraph_format.left_indent = Cm(0)
        toc1.paragraph_format.space_before = Pt(6)
        toc1.paragraph_format.space_after = Pt(3)

        # TOC 2 - 二级目录
        try:
            toc2 = styles["TOC 2"]
        except KeyError:
            toc2 = styles.add_style("TOC 2", 1)

        toc2.font.name = "宋体"
        toc2.font.size = Pt(12)
        toc2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        toc2.paragraph_format.left_indent = Cm(0.74)  # 约2字符
        toc2.paragraph_format.space_before = Pt(3)
        toc2.paragraph_format.space_after = Pt(0)

        # TOC 3 - 三级目录
        try:
            toc3 = styles["TOC 3"]
        except KeyError:
            toc3 = styles.add_style("TOC 3", 1)

        toc3.font.name = "宋体"
        toc3.font.size = Pt(12)
        toc3._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        toc3.paragraph_format.left_indent = Cm(1.48)  # 约4字符
        toc3.paragraph_format.space_before = Pt(0)
        toc3.paragraph_format.space_after = Pt(0)

        logger.info("目录样式设置完成")

    def setup_heading_styles(self, format_settings=None):
        """
        设置标题样式

        Args:
            format_settings: 格式设置字典
        """
        if format_settings is None:
            format_settings = {}

        styles = self.doc.styles

        for level in range(1, 5):
            style_name = f"Heading {level}"
            if style_name in styles:
                style = styles[style_name]

                font_key = f"heading{level}_font"
                size_key = f"heading{level}_size"
                bold_key = f"heading{level}_bold"

                font_name = format_settings.get(font_key, "黑体")
                font_size = format_settings.get(size_key, 16 - (level - 1) * 2)
                is_bold = format_settings.get(bold_key, True)

                style.font.name = font_name
                style.font.size = Pt(font_size)
                style.font.bold = is_bold
                style.font.italic = False  # 🔑 确保不斜体
                style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

                # 🔑 关键修复：显式设置样式颜色为黑色，清除Word内置的蓝色
                from docx.shared import RGBColor

                style.font.color.rgb = RGBColor(0, 0, 0)

                # 🔑 在XML层面也确保颜色为黑色
                rPr = style._element.get_or_add_rPr()
                color_elem = rPr.find(qn("w:color"))
                if color_elem is not None:
                    rPr.remove(color_elem)
                color = OxmlElement("w:color")
                color.set(qn("w:val"), "000000")
                rPr.append(color)

                # 设置段落格式
                if level == 1:
                    style.paragraph_format.space_before = Pt(24)
                    style.paragraph_format.space_after = Pt(12)
                elif level == 2:
                    style.paragraph_format.space_before = Pt(12)
                    style.paragraph_format.space_after = Pt(6)
                else:
                    style.paragraph_format.space_before = Pt(6)
                    style.paragraph_format.space_after = Pt(3)

        logger.info("标题样式设置完成")

    def setup_normal_style(self, format_settings=None):
        """
        设置正文样式

        Args:
            format_settings: 格式设置字典
        """
        if format_settings is None:
            format_settings = {}

        style = self.doc.styles["Normal"]

        body_font = format_settings.get("body_font", "宋体")
        body_size = format_settings.get("body_size", 12)
        line_spacing = format_settings.get("body_line_spacing", 1.5)

        style.font.name = body_font
        style.font.size = Pt(body_size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)

        style.paragraph_format.line_spacing = line_spacing

        logger.info("正文样式设置完成")
