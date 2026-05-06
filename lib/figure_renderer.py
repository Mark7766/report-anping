# -*- coding: utf-8 -*-
"""
ajepro backend - 图片渲染引擎
处理图片的专业化渲染，包括自动编号、图题、多图并排等
"""

import os
import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# 使用数值常量代替枚举
WD_ALIGN_PARAGRAPH_CENTER = 1
WD_ALIGN_PARAGRAPH_LEFT = 0
WD_ALIGN_PARAGRAPH_RIGHT = 2

from lib.logger import get_logger

logger = get_logger("figure_renderer")


class FigureNumbering:
    """图片编号追踪器"""

    def __init__(self):
        self.chapter = 0
        self.figure_count = 0
        self.figure_registry = {}  # id -> number mapping

    def set_chapter(self, chapter_num):
        """设置当前章节号（支持数字和字符串，如"前"、"附A"）"""
        # 🔑 支持字符串章节号（如"前"、"附A"）
        if isinstance(chapter_num, str):
            # 字符串章节号，直接使用
            if chapter_num != self.chapter:
                self.chapter = chapter_num
                self.figure_count = 0
        else:
            # 数字章节号，进行保护性检查
            if chapter_num <= 0:
                chapter_num = 4
                logger.warning(f"[FigureNumbering] 检测到无效章节号，使用默认值: {chapter_num}")

            if chapter_num != self.chapter:
                self.chapter = chapter_num
                self.figure_count = 0

    def next_number(self, figure_id=None):
        """
        获取下一个图片编号

        Args:
            figure_id: 可选的图片ID用于交叉引用

        Returns:
            str: 图片编号，如 "3-1", "前-1", "附A-1"
        """
        # 🔑 保护性检查：确保章节号有效（仅对数字章节）
        if isinstance(self.chapter, int) and self.chapter <= 0:
            self.chapter = 4
            logger.warning(f"[FigureNumbering] 图片编号时检测到无效章节号，使用默认值: {self.chapter}")

        self.figure_count += 1
        number = f"{self.chapter}-{self.figure_count}"

        if figure_id:
            self.figure_registry[figure_id] = number

        return number

    def get_number(self, figure_id):
        """根据ID获取图片编号"""
        return self.figure_registry.get(figure_id)


class ChapterNumberingTracker:
    """章节编号追踪器 - 用于追踪当前处于哪个章节"""

    def __init__(self):
        self.current_chapter = 0
        self.chapter_pattern = re.compile(r"^(?:第\s*(\d+)\s*章|(\d+)[\.\s])")
        # 🔑 特殊章节标识符映射
        self.special_chapters = {
            "前言": "前",
            "序言": "序",
            "绪论": "绪",
            "摘要": "摘",
            "附录a": "附A",
            "附录b": "附B",
            "附录c": "附C",
            "附录d": "附D",
            "附录e": "附E",
        }
        self.current_chapter_text = None  # 🔑 存储当前章节的文本标识（如"前"、"附A"）

    def detect_initial_chapter(self, content):
        """
        检测内容中的第一个章节号

        Args:
            content: Markdown内容

        Returns:
            int: 检测到的章节号，如果没有检测到返回4（默认值）
        """
        lines = content.split("\n") if isinstance(content, str) else content

        for line in lines:
            # 只处理一级标题
            if line.startswith("#") and not line.startswith("##"):
                heading_text = line[1:].strip()

                # 🔑 先检查是否是特殊章节
                heading_lower = heading_text.lower()
                for key, value in self.special_chapters.items():
                    if key in heading_lower:
                        self.current_chapter_text = value
                        return 0  # 特殊章节返回0，但会设置chapter_text

                # 检查数字章节
                match = self.chapter_pattern.match(heading_text)
                if match:
                    chapter_num = match.group(1) or match.group(2)
                    try:
                        detected = int(chapter_num)
                        if detected > 0:  # 确保章节号大于0
                            self.current_chapter_text = None
                            return detected
                    except ValueError:
                        pass

        # 如果没有检测到有效章节号，返回默认值4（通常是"场地地震动参数"章节）
        self.current_chapter_text = None
        return 4

    def initialize_from_content(self, content):
        """
        从内容初始化章节号

        Args:
            content: Markdown内容
        """
        self.current_chapter = self.detect_initial_chapter(content)

    def update_from_heading(self, heading_text, level):
        """
        从标题文本更新当前章节号

        Args:
            heading_text: 标题文本
            level: 标题级别 (1-4)
        """
        if level == 1:
            # 🔑 先检查是否是特殊章节
            heading_lower = heading_text.strip().lower()
            for key, value in self.special_chapters.items():
                if key in heading_lower:
                    self.current_chapter = 0
                    self.current_chapter_text = value
                    return

            # 尝试提取数字章节号
            match = self.chapter_pattern.match(heading_text.strip())
            if match:
                chapter_num = match.group(1) or match.group(2)
                try:
                    self.current_chapter = int(chapter_num)
                    self.current_chapter_text = None
                except ValueError:
                    self.current_chapter += 1
                    self.current_chapter_text = None
            else:
                self.current_chapter += 1
                self.current_chapter_text = None

    def get_current_chapter(self):
        """获取当前章节号或章节文本标识"""
        # 🔑 如果有特殊章节文本，返回它
        if self.current_chapter_text:
            return self.current_chapter_text
        # 保护性检查：如果章节号为0且没有特殊文本，返回默认值4
        if self.current_chapter <= 0:
            return 4
        return self.current_chapter


class DocxFigureRenderer:
    """Word图片渲染器"""

    def __init__(self, doc, format_settings=None, figure_numbering=None):
        """
        初始化图片渲染器

        Args:
            doc: Word文档对象
            format_settings: 格式设置字典
            figure_numbering: 图片编号追踪器
        """
        self.doc = doc
        self.format_settings = format_settings or {}
        self.numbering = figure_numbering or FigureNumbering()

    def render(
        self,
        img_path,
        caption=None,
        figure_id=None,
        specified_width=None,
        specified_height=None,
        add_number=True,
        preset_number=None,
    ):
        """
        渲染图片到Word文档

        Args:
            img_path: 图片文件路径
            caption: 图片标题（如果已包含编号如"图 3-1 xxx"，则设置add_number=False）
            figure_id: 图片ID（用于交叉引用）
            specified_width: 指定宽度（像素）
            specified_height: 指定高度（像素）
            add_number: 是否添加自动编号
            preset_number: 预设的图片编号（如 "3-1"）

        Returns:
            bool: 是否成功
        """
        try:
            # 检查文件是否存在
            if not os.path.exists(img_path):
                self._add_placeholder(caption or img_path, "图片未找到")
                # 🔑 即使图片不存在，也添加图题（带编号）
                self._add_figure_caption(caption, figure_id, add_number, preset_number)
                logger.warning(f"图片文件不存在: {img_path}")
                return False

            # 计算图片尺寸
            final_width_cm = self._calculate_image_width(img_path, specified_width, specified_height)

            # 添加图片段落
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH_CENTER

            # 添加图片
            run = p.add_run()
            try:
                run.add_picture(img_path, width=Cm(final_width_cm))
            except Exception as pic_err:
                logger.warning(f"添加图片失败 {img_path}: {pic_err}")
                self._add_placeholder(caption or img_path, "图片加载失败")
                # 🔑 图片加载失败也添加图题
                self._add_figure_caption(caption, figure_id, add_number, preset_number)
                return False

            # 添加图题
            self._add_figure_caption(caption, figure_id, add_number, preset_number)

            return True

        except Exception as e:
            logger.error(f"渲染图片失败 {img_path}: {str(e)}")
            self._add_placeholder(caption or img_path, "图片处理失败")
            # 🔑 异常情况也添加图题
            self._add_figure_caption(caption, figure_id, add_number, preset_number)
            return False

    def render_group(self, images, columns=2, group_caption=None):
        """
        渲染多图并排布局

        Args:
            images: 图片列表 [{'path': str, 'caption': str}, ...]
            columns: 列数
            group_caption: 组图标题
        """
        if not images:
            return

        # 计算行数
        rows = (len(images) + columns - 1) // columns

        # 创建无边框表格
        table = self.doc.add_table(rows=rows, cols=columns)
        table.alignment = 1  # 居中

        # 移除边框
        self._remove_table_borders(table)

        # 计算每列图片宽度
        margin_left = self.format_settings.get("page_margin_left", 3.0)
        margin_right = self.format_settings.get("page_margin_right", 2.5)
        available_width = 21.0 - margin_left - margin_right
        cell_width = (available_width - 0.5) / columns  # 留一点间隙

        # 填充图片
        for idx, img_info in enumerate(images):
            row_idx = idx // columns
            col_idx = idx % columns

            cell = table.rows[row_idx].cells[col_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH_CENTER

            img_path = img_info.get("path", "")
            sub_caption = img_info.get("caption", "")
            if os.path.exists(img_path):
                try:
                    run = p.add_run()
                    run.add_picture(img_path, width=Cm(cell_width * 0.9))

                    # 添加子图标题
                    if sub_caption:
                        cap_p = cell.add_paragraph()
                        cap_p.alignment = WD_ALIGN_PARAGRAPH_CENTER
                        cap_run = cap_p.add_run(f"({chr(97 + idx)}) {sub_caption}")
                        cap_run.font.size = Pt(9)
                        cap_run.font.name = "宋体"
                        cap_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                except Exception as e:
                    logger.warning(f"添加组图图片失败: {e}")
                    run = p.add_run(f"[图片: {sub_caption}]")

        # 添加组图标题
        if group_caption:
            figure_number = self.numbering.next_number()
            self._add_figure_caption(group_caption, None, True, figure_number)

    def _calculate_image_width(self, img_path, specified_width=None, specified_height=None):
        """计算图片最终宽度"""
        margin_left = self.format_settings.get("page_margin_left", 3.0)
        margin_right = self.format_settings.get("page_margin_right", 2.5)
        available_width_cm = 21.0 - margin_left - margin_right

        final_width_cm = available_width_cm

        try:
            from PIL import Image

            with Image.open(img_path) as img:
                img_width_px, img_height_px = img.size
                img_width_cm = img_width_px / 96 * 2.54

                if specified_width:
                    final_width_cm = min(int(specified_width) / 96 * 2.54, available_width_cm)
                elif img_width_cm <= available_width_cm:
                    # 小图片不放大
                    final_width_cm = img_width_cm
                else:
                    # 大图片缩放
                    final_width_cm = available_width_cm
        except Exception as e:
            logger.warning(f"无法读取图片尺寸 {img_path}: {e}")
            final_width_cm = min(15, available_width_cm)

        return final_width_cm

    def _add_figure_caption(self, caption, figure_id=None, add_number=True, preset_number=None):
        """添加图题"""
        caption_font = self.format_settings.get("figure_caption_font", "黑体")
        caption_size = self.format_settings.get("figure_caption_size", 10.5)

        # 获取编号
        if add_number:
            if preset_number:
                figure_number = preset_number
            else:
                figure_number = self.numbering.next_number(figure_id)

            caption_text = f"图{figure_number}"
            if caption:
                caption_text += f"  {caption}"
        else:
            caption_text = caption or ""

        if not caption_text:
            return

        # 添加图题段落
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH_CENTER

        run = p.add_run(caption_text)
        run.font.name = caption_font
        run.font.size = Pt(caption_size)
        run.font.bold = True  # 🔑 加粗
        run.font.italic = False  # 🔑 不使用斜体
        from docx.shared import RGBColor

        run.font.color.rgb = RGBColor(0, 0, 0)  # 🔑 显式设置为黑色
        run._element.rPr.rFonts.set(qn("w:eastAsia"), caption_font)

        # 设置间距
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        # 设置段后间距
        p.paragraph_format.space_after = Pt(12)

    def _add_placeholder(self, caption, error_msg):
        """添加图片占位符"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH_CENTER
        run = p.add_run(f"[{error_msg}: {caption}]")
        run.font.name = "宋体"
        run.font.size = Pt(10)
        run.italic = True
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def _remove_table_borders(self, table):
        """移除表格边框"""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

        tblBorders = OxmlElement("w:tblBorders")
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")
            tblBorders.append(border)

        tblPr.append(tblBorders)

        if tbl.tblPr is None:
            tbl.insert(0, tblPr)
