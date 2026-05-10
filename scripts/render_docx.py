from __future__ import annotations

"""
scripts/render_docx.py — 将 chapters/*.md 和 params.json 渲染为 .docx 报告

调用链:
  argparse → load_params / load_chapters →
  MarkdownRenderer（解析 MD 内容）→
  DocxBuilder / DocxTableRenderer / DocxFigureRenderer →
  Document.save()

用法:
    python scripts/render_docx.py \\
        --params params_example.json \\
        --chapters chapters/ \\
        --out exports/report.docx
"""

import argparse
import json
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt

from lib.docx_builder import DocxBuilder
from lib.figure_renderer import (
    ChapterNumberingTracker,
    DocxFigureRenderer,
    FigureNumbering,
)
from lib.logger import get_logger
from lib.table_renderer import (
    DocxTableRenderer,
    MarkdownTableParser,
    TableNumbering,
)

logger = get_logger("render_docx")

# ============================================================
# 章节文件排序键：NN_xxx.md → (sort_key, path)
# ============================================================

_CHAPTER_ORDER = [
    "preface",
    "chapter1",
    "chapter2",
    "chapter3",
    "chapter4",
    "chapter5",
    "chapter6",
    "chapter7",
    "chapter8",
    "chapter9",
    "chapter10",
    "appendix",
]


def _sort_key(path: Path) -> tuple[int, str]:
    """Return (priority, stem) so chapters are in GB 17741-2025 order."""
    stem = path.stem.lower()
    # Strip leading digits: "01_chapter1" → "chapter1"
    clean = re.sub(r"^\d+[_\-]?", "", stem)
    try:
        idx = _CHAPTER_ORDER.index(clean)
    except ValueError:
        idx = len(_CHAPTER_ORDER)
    return (idx, stem)


# ============================================================
# Markdown → Docx renderer
# ============================================================


class MarkdownToDocxRenderer:
    """
    简单的 Markdown → Word 渲染器。
    支持: #/##/###/#### 标题、段落、表格、图片引用（![...](...)）、
    内联 **粗体** / *斜体* 格式。
    """

    # 图片引用: ![caption](path) 或 ![caption](path){width=N}
    _IMG_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)(?:\{[^}]*\})?")
    # 表格行: 以 | 开头
    _TABLE_LINE_RE = re.compile(r"^\s*\|")
    # 内联格式分割: **粗体** 或 *斜体*（先匹配双星号）
    _INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    # 图题已含编号: 以 "图N-N" 开头
    _CAPTION_HAS_NUM_RE = re.compile(r"^图\s*\d+[-\u2013]\d+")

    def __init__(
        self,
        doc: Document,
        format_settings: dict | None = None,
        base_dir: Path | None = None,
    ) -> None:
        """Initialise with a python-docx Document and optional format_settings."""
        self.doc = doc
        self.fmt = format_settings or {}
        self.base_dir = base_dir  # project root for resolving relative image paths
        self.table_renderer = DocxTableRenderer(doc, self.fmt)
        self.table_numbering = TableNumbering()
        self.figure_numbering = FigureNumbering()
        self.figure_renderer = DocxFigureRenderer(doc, self.fmt, self.figure_numbering)
        self.chapter_tracker = ChapterNumberingTracker()

    def _add_paragraph_with_inline_fmt(self, text: str) -> None:
        """Add a paragraph to doc with **bold** and *italic* inline markdown parsed into runs."""
        para = self.doc.add_paragraph()
        para.alignment = 3  # justify
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.first_line_indent = Pt(24)
        para.paragraph_format.space_after = Pt(6)
        body_font = self.fmt.get("body_font", "宋体")
        body_size = self.fmt.get("body_size", 12)

        parts = self._INLINE_RE.split(text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                run.italic = True
            else:
                run = para.add_run(part)
            run.font.name = body_font
            run.font.size = Pt(body_size)

    def render_chapter(self, content: str) -> None:
        """Render a complete chapter's markdown content into self.doc."""
        self.chapter_tracker.initialize_from_content(content)
        ch = self.chapter_tracker.get_current_chapter()
        self.table_numbering.set_chapter(ch)
        self.figure_numbering.set_chapter(ch)

        lines = content.splitlines()
        idx = 0
        while idx < len(lines):
            line = lines[idx]

            # --- Heading ---
            heading_match = re.match(r"^(#{1,4})\s+(.*)", line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                self.doc.add_heading(text, level=level)
                self.chapter_tracker.update_from_heading(text, level)
                ch = self.chapter_tracker.get_current_chapter()
                self.table_numbering.set_chapter(ch)
                self.figure_numbering.set_chapter(ch)
                idx += 1
                continue

            # --- Image (check BEFORE table so | ![...](...) | lines are caught) ---
            img_match = self._IMG_RE.search(line)
            if img_match:
                caption = img_match.group(1).strip()
                img_path = img_match.group(2).strip()
                # Resolve relative paths against the project root (base_dir) so that
                # paths like "assets/generated/xxx.png" work regardless of CWD.
                if self.base_dir and not Path(img_path).is_absolute():
                    img_path = str(self.base_dir / img_path)
                # If caption already contains a figure number (e.g. "图2-1 ..."),
                # skip auto-numbering to avoid double numbers like "图2-1  图2-1 ...".
                add_num = not bool(self._CAPTION_HAS_NUM_RE.match(caption))
                self.figure_renderer.render(img_path, caption=caption, add_number=add_num)
                idx += 1
                continue

            # --- Table (collect consecutive | lines) ---
            if self._TABLE_LINE_RE.match(line):
                table_lines: list[str] = []
                while idx < len(lines) and self._TABLE_LINE_RE.match(lines[idx]):
                    table_lines.append(lines[idx])
                    idx += 1
                table_data = MarkdownTableParser.parse(table_lines)
                if table_data:
                    self.table_renderer.render(
                        table_data,
                        table_number=self.table_numbering.next_number(),
                    )
                continue

            # --- Blank line or paragraph text ---
            stripped = line.strip()
            if stripped:
                self._add_paragraph_with_inline_fmt(stripped)
            idx += 1


# ============================================================
# Core render function
# ============================================================


def render(
    params: dict,
    chapters_dir: Path,
    out_path: Path,
) -> None:
    """
    渲染完整报告 .docx。

    Args:
        params: 项目参数字典（来自 params.json）
        chapters_dir: 包含章节 .md 文件的目录
        out_path: 输出 .docx 路径
    """
    md_files = sorted(chapters_dir.glob("*.md"), key=_sort_key)
    if not md_files:
        logger.warning(f"chapters/ 目录中没有 .md 文件: {chapters_dir}")

    doc = Document()
    builder = DocxBuilder(doc)

    # 页面设置
    builder.setup_page()
    builder.setup_heading_styles()
    builder.setup_normal_style()

    # 封面信息
    project_name = params.get("name", "地震安全性评价报告")
    title_para = doc.add_paragraph()
    title_para.alignment = 1  # center
    run = title_para.add_run(project_name)
    run.font.bold = True
    run.font.size = Pt(22)
    doc.add_paragraph(f"工作等级: {params.get('level', 'II')} 级")
    doc.add_paragraph(f"评价单位: {params.get('evaluation_unit', '')}")
    doc.add_paragraph(f"报告日期: {params.get('report_date', '')}")
    doc.add_page_break()

    # 目录（Word 原生可更新域）
    builder.add_toc()

    # 页眉页脚
    builder.setup_headers_footers(
        header_odd=project_name,
        header_even="地震安全性评价报告",
    )
    builder.setup_toc_styles()

    # 渲染各章节
    fmt_settings: dict = {}
    md_renderer = MarkdownToDocxRenderer(doc, fmt_settings, base_dir=chapters_dir.resolve().parent)
    for md_file in md_files:
        content = md_file.read_text(encoding="utf-8")
        md_renderer.render_chapter(content)
        logger.info(f"渲染章节: {md_file.name}")

    # 保存
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    logger.info(f"报告已保存: {out_path}")
    print(f"[OK] 报告已生成: {out_path}  ({out_path.stat().st_size:,} bytes)")


# ============================================================
# CLI 入口
# ============================================================


def main() -> None:
    """CLI 入口。"""
    parser = argparse.ArgumentParser(description="将 chapters/*.md 渲染为 GB 17741-2025 安评报告 .docx")
    parser.add_argument(
        "--params",
        default="params.json",
        help="项目参数 JSON 文件路径，默认 params.json",
    )
    parser.add_argument(
        "--chapters",
        default="chapters/",
        help="章节 .md 文件目录，默认 chapters/",
    )
    parser.add_argument(
        "--out",
        default="exports/report.docx",
        help="输出 .docx 文件路径，默认 exports/report.docx",
    )
    args = parser.parse_args()

    params_path = Path(args.params)
    if not params_path.exists():
        print(f"[ERROR] 参数文件不存在: {args.params}", file=sys.stderr)
        sys.exit(1)
    params = json.loads(params_path.read_text(encoding="utf-8"))

    chapters_dir = Path(args.chapters)
    if not chapters_dir.is_dir():
        print(f"[ERROR] 章节目录不存在: {args.chapters}", file=sys.stderr)
        sys.exit(1)

    render(params=params, chapters_dir=chapters_dir, out_path=Path(args.out))


if __name__ == "__main__":
    main()
